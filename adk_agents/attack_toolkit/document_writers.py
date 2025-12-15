from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from fpdf import FPDF
from .config import PDF_SETTINGS, DOCX_SETTINGS

class DocxWriter:
    """Creates a DOCX file with a hidden payload."""
    
    def create_document(self, payload: str, visible_text: str, file_path: str = "attack.docx"):
        doc = Document()
        doc.add_heading(DOCX_SETTINGS.get("default_heading", "Document"), 0)
        
        # Add the visible, benign text
        doc.add_paragraph(visible_text)
        
        # Add the hidden payload
        p = doc.add_paragraph()
        run = p.add_run(payload)
        font = run.font
        
        if DOCX_SETTINGS.get("use_hidden_text", True):
            font.hidden = True
        
        if DOCX_SETTINGS.get("use_white_text", False):
            font.color.rgb = (255, 255, 255)
            
        doc.save(file_path)
        return file_path

class PdfWriter:
    """Creates a PDF file with a hidden payload."""
    
    def create_document(self, payload: str, visible_text: str, file_path: str = "attack.pdf"):
        pdf = FPDF()
        pdf.add_page()
        
        # Add the visible, benign text
        pdf.set_font(PDF_SETTINGS.get("default_font"), size=PDF_SETTINGS.get("default_size"))
        pdf.multi_cell(0, 5, visible_text)
        pdf.ln(10) # Add a line break
        
        # Add the hidden payload
        color = PDF_SETTINGS.get("hidden_text_color", (255, 255, 255))
        pdf.set_text_color(color[0], color[1], color[2])
        
        # Try to use a font that supports Unicode for the payload
        try:
            # Try a standard Unicode font
            pdf.set_font("DejaVu", size=PDF_SETTINGS.get("default_size"))
        except RuntimeError:
            # Fallback to a specified font path or default
            font_path = PDF_SETTINGS.get("unicode_font_path")
            if font_path:
                try:
                    pdf.add_font("CustomUnicode", "", font_path, uni=True)
                    pdf.set_font("CustomUnicode", size=PDF_SETTINGS.get("default_size"))
                except Exception as e:
                    print(f"Warning: Could not load unicode font. Payload may not render correctly. {e}")
                    pdf.set_font(PDF_SETTINGS.get("default_font"), size=PDF_SETTINGS.get("default_size"))
            else:
                 pdf.set_font(PDF_SETTINGS.get("default_font"), size=PDF_SETTINGS.get("default_size"))
        
        pdf.multi_cell(0, 5, payload)
        
        pdf.output(file_path)
        return file_path
