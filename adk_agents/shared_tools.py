import os
import uuid
from docx import Document
from docx.shared import RGBColor

# Import ReportLab for robust PDF generation
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    print("Warning: 'reportlab' not installed. PDF generation will fail. Run: uv pip install reportlab")
    REPORTLAB_AVAILABLE = False

# Import toolkit
try:
    from attack_toolkit.leet_attack import LeetAttack
    from attack_toolkit.homoglyph_attack import HomoglyphAttack
    from attack_toolkit.splitting_attack import PromptSplittingAttack
    from attack_toolkit.unicode_stega_attack import UnicodeStegaAttack
    from attack_toolkit.whitespace_attack import WEncodingAttack
    from attack_toolkit.token_splitting_attack import TokenSplittingAttack
except ImportError:
    print("Warning: Shared Tools could not import attack_toolkit.")

class DocxWriter:
    def create_document(self, payload: str, visible_text: str, file_path: str):
        doc = Document()
        doc.add_heading("Internal Document", 0)
        doc.add_paragraph(visible_text)
        
        # Add payload (DOCX handles Unicode natively)
        p = doc.add_paragraph()
        run = p.add_run(payload)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        doc.save(file_path)
        return file_path

class PdfWriter:
    def __init__(self):
        self.font_name = "Helvetica" # Default standard font
        self.unicode_support = False
        if REPORTLAB_AVAILABLE:
            self._register_unicode_font()

    def _register_unicode_font(self):
        """Locates and registers a usable Unicode font for ReportLab."""
        # Common paths for fonts in Linux/WSL
        search_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
            # Add a local fallback if you drop a ttf in the folder
            os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf") 
        ]
        
        for path in search_paths:
            if os.path.exists(path):
                try:
                    # Register the font with ReportLab
                    pdfmetrics.registerFont(TTFont('UnicodeFont', path))
                    self.font_name = 'UnicodeFont'
                    self.unicode_support = True
                    print(f"--- [SharedTool] Registered Unicode font: {path} ---")
                    return
                except Exception as e:
                    print(f"Warning: Failed to register font {path}: {e}")
        
        print("Warning: No Unicode font found. PDF generation for homoglyphs may fail.")

    def create_document(self, payload: str, visible_text: str, file_path: str):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab library is not installed")

        c = canvas.Canvas(file_path, pagesize=letter)
        
        # Setup Text Object
        text_object = c.beginText(40, 750)
        text_object.setFont(self.font_name, 12)
        
        # Add Visible Text (Black)
        c.setFillColorRGB(0, 0, 0)
        for line in visible_text.split('\n'):
            text_object.textLine(line)
        
        text_object.textLine("") # Spacing
        
        # Add Payload (White/Invisible)
        # We reset the text object to ensure color changes apply correctly
        c.drawText(text_object)
        
        payload_object = c.beginText(40, 700) # Position below visible text
        payload_object.setFont(self.font_name, 12)
        payload_object.setFillColorRGB(1, 1, 1) # White text
        
        try:
            for line in payload.split('\n'):
                payload_object.textLine(line)
            c.drawText(payload_object)
        except Exception as e:
            print(f"--- [SharedTool] PDF Write Error (likely font missing glyphs): {e} ---")
            c.setFillColorRGB(1, 0, 0) # Red warning
            c.drawString(40, 680, "[PAYLOAD ENCODING FAILED - SEE LOGS]")

        c.save()
        return file_path

def create_malicious_file(payload: str, encoding: str = "none", file_type: str = "docx") -> str:
    """
    Creates a file (PDF/DOCX) containing the payload.
    AUTO-FALLBACK: Switches to DOCX if PDF generation fails for ANY reason.
    """
    print(f"--- [SharedTool] Creating {file_type} with {encoding} ---")
    file_type = file_type.lower()
    
    encoded_payload = payload
    
    # Apply Encoding
    try:
        if encoding == "leet": encoded_payload = LeetAttack(strength=1.0).transform(payload)
        elif encoding == "homoglyph": encoded_payload = HomoglyphAttack(strength=1.0).transform(payload)
        elif encoding == "splitting": encoded_payload = PromptSplittingAttack(strength=0.7).transform(payload)
        elif encoding == "unicode_stega": encoded_payload = UnicodeStegaAttack(strength=1.0).transform(payload)
        elif encoding == "whitespace": encoded_payload = WEncodingAttack().transform(payload)
        elif encoding == "token_split": encoded_payload = TokenSplittingAttack().transform(payload)
    except NameError:
        return "Error: Attack Toolkit modules not loaded."

    visible_text = "CONFIDENTIAL: Internal Audit Report 2025. Do not distribute."
    
    output_dir = "generated_attacks"
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"attack_{uuid.uuid4().hex[:8]}"

    try:
        if file_type == "pdf":
            try:
                writer = PdfWriter()
                file_path = os.path.join(output_dir, f"{filename}.pdf")
                writer.create_document(encoded_payload, visible_text, file_path)
                return file_path
            except Exception as e:
                print(f"--- [SharedTool] PDF Creation Failed ({e}). Falling back to DOCX. ---")
                # Fall through to DOCX creation below
                file_type = "docx"

        # DOCX Creation (Default or Fallback)
        writer = DocxWriter()
        file_path = os.path.join(output_dir, f"{filename}.docx")
        writer.create_document(encoded_payload, visible_text, file_path)
        
        return file_path

    except Exception as e:
        return f"Error creating file: {e}"